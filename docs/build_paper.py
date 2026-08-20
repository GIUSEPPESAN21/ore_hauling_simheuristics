"""Regenerate PAPER_ore_hauling_simheuristics.docx and the HTML artifact from
PAPER_ore_hauling_simheuristics.md. Run this after editing the manuscript source
(e.g. once the author team fills in the bracketed front-matter placeholders).

Usage: python docs/build_paper.py [docx|html|all]   (default: all)

Requires python-docx and matplotlib in the active environment:
    pip install python-docx matplotlib
(not in requirements.txt: these are manuscript-tooling dependencies, not part
of the simulation/optimization codebase itself.)
"""
from __future__ import annotations

import base64
import re
import sys
from pathlib import Path

DOCS = Path(__file__).resolve().parent
PROJECT = DOCS.parent
SRC = DOCS / "PAPER_ore_hauling_simheuristics.md"
FIGDIR = DOCS / "figures"
DOCX_OUT = DOCS / "PAPER_ore_hauling_simheuristics.docx"
HTML_OUT = DOCS / "paper_ore_hauling_simheuristics_artifact.html"

FIGURE_FILES = {
    "Figure 1.": "figure1_utilization.png",
    "Figure 2.": "figure2_convergence.png",
}

# Markdown backslash-escapes (`\*`, `\_`, `\|`, ...) must survive as literal
# characters without being read as emphasis/table-delimiter markup. Protect
# them with visible-but-unlikely bracket tokens before tokenizing bold/italic/
# backtick spans, then restore the literal character in the final run text.
_ESCAPES = {"\\*": "", "\\_": "", "\\|": "", "\\`": "",
            "\\[": "", "\\]": ""}
_UNESCAPES = {v: k[1] for k, v in _ESCAPES.items()}


def protect_escapes(text):
    for esc, sentinel in _ESCAPES.items():
        text = text.replace(esc, sentinel)
    return text


def restore_escapes(text):
    for sentinel, lit in _UNESCAPES.items():
        text = text.replace(sentinel, lit)
    return text


# --------------------------------------------------------------------------- DOCX
# Visual target: the MDPI journal-article template (template.pdf) --
# italic "Article" eyebrow; bold title; superscripted author/affiliation block;
# a page-1 sidebar (Received/Revised/Accepted/Published + CC-BY copyright notice)
# beside the Abstract; sans-serif bold numbered headings over serif body text;
# three-line ("booktabs") tables with no vertical rules; back-matter sections
# (Author Contributions, Funding, ...) as a single run-in-bold-label paragraph,
# not a separate heading; a running header "Version <date> submitted to ... | N of M"
# on every page except the first; a DOI-style footer on every page.

BODY_FONT = "Cambria"       # serif body face, close in feel to MDPI's Palatino-style body
HEAD_FONT = "Calibri"       # bold sans-serif used for numbered section headings
JOURNAL_LINE = "Version August 19, 2026 submitted to Journal Not Specified"
DOI_LINE = "https://doi.org/10.3390/xxxxxxx"

BACK_MATTER_LABELS = (
    "Author Contributions:", "Funding:", "Institutional Review Board Statement:",
    "Informed Consent Statement:", "Data Availability Statement:", "Acknowledgments:",
    "Conflicts of Interest:",
)


def build_docx():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    inline_sup = re.compile(r"<sup>(.*?)</sup>")
    inline_sub = re.compile(r"<sub>(.*?)</sub>")
    INK = RGBColor(0x15, 0x15, 0x15)
    MUTED = RGBColor(0x55, 0x55, 0x55)
    CODE_FONT = "Consolas"

    def add_inline_runs(paragraph, text, base_size=10.5, base_bold=False, base_italic=False,
                         font=BODY_FONT, color=None):
        text = protect_escapes(text)
        token_re = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|<sup>.*?</sup>|<sub>.*?</sub>)")
        for part in token_re.split(text):
            if not part:
                continue
            code_run = False
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(restore_escapes(part[2:-2])); run.bold = True
            elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
                run = paragraph.add_run(restore_escapes(part[1:-1])); code_run = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = paragraph.add_run(restore_escapes(part[1:-1])); run.italic = True
            elif part.startswith("<sup>"):
                run = paragraph.add_run(restore_escapes(inline_sup.match(part).group(1))); run.font.superscript = True
            elif part.startswith("<sub>"):
                run = paragraph.add_run(restore_escapes(inline_sub.match(part).group(1))); run.font.subscript = True
            else:
                run = paragraph.add_run(restore_escapes(part)); run.bold = base_bold; run.italic = base_italic
            run.font.size = Pt(base_size * 0.93 if code_run else base_size)
            run.font.name = CODE_FONT if code_run else font
            if color is not None:
                run.font.color.rgb = color

    def set_cell_border(cell, **edges):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge, spec in edges.items():
            el = OxmlElement(f'w:{edge}')
            for k, v in spec.items():
                el.set(qn(f'w:{k}'), str(v))
            borders.append(el)
        tcPr.append(borders)

    def repeat_header_row(row):
        """Sets w:tblHeader so this row reprints atop every page the table spans."""
        trPr = row._tr.get_or_add_trPr()
        el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true')
        trPr.append(el)

    def three_line_table(table):
        """Booktabs style: rule above the header, rule below the header, rule
        below the last row; no vertical rules, no rules between body rows."""
        thick = {"sz": 12, "val": "single", "color": "151515"}
        thin = {"sz": 6, "val": "single", "color": "151515"}
        nrows = len(table.rows)
        repeat_header_row(table.rows[0])
        for cidx, cell in enumerate(table.rows[0].cells):
            set_cell_border(cell, top=thick, bottom=thin)
        if nrows > 1:
            for cell in table.rows[-1].cells:
                set_cell_border(cell, bottom=thick)

    def set_no_borders(table):
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top={"sz": 0, "val": "nil"}, bottom={"sz": 0, "val": "nil"},
                                 left={"sz": 0, "val": "nil"}, right={"sz": 0, "val": "nil"})

    def add_table(doc, lines):
        rows = [l for l in lines if l.strip().startswith("|")]
        data_rows = []
        for idx, r in enumerate(rows):
            cells = [c.strip() for c in r.strip().strip("|").split("|")]
            if idx == 1 and all(re.match(r"^:?-{2,}:?$", c) for c in cells):
                continue
            data_rows.append(cells)
        if not data_rows:
            return
        ncols = len(data_rows[0])
        table = doc.add_table(rows=0, cols=ncols)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ridx, row in enumerate(data_rows):
            cells = table.add_row().cells
            for cidx in range(ncols):
                text = row[cidx] if cidx < len(row) else ""
                p = cells[cidx].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                add_inline_runs(p, text, base_size=8.5, base_bold=(ridx == 0), font=HEAD_FONT if ridx == 0 else BODY_FONT)
        three_line_table(table)
        cap = doc.add_paragraph(); cap.paragraph_format.space_after = Pt(14)

    def add_figure(doc, key):
        path = FIGDIR / FIGURE_FILES[key]
        if path.exists():
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.add_run().add_picture(str(path), width=Cm(14.5))

    def add_code_block(doc, code_lines):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run("\n".join(code_lines)); run.font.name = "Consolas"; run.font.size = Pt(8.5)

    def add_field(paragraph, instr_text, fallback="1"):
        def _r():
            return paragraph.add_run()._r
        f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); _r().append(f1)
        instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = instr_text
        _r().append(instr)
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate'); _r().append(f2)
        t = OxmlElement('w:t'); t.text = fallback; _r().append(t)
        f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end'); _r().append(f3)

    def heading_paragraph(doc, text, size, space_before=18, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        add_inline_runs(p, text, base_size=size, base_bold=True, font=HEAD_FONT, color=INK)
        return p

    # ---- document + page setup -------------------------------------------------
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.12
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.4)
    section.top_margin, section.bottom_margin = Cm(2.6), Cm(2.2)
    section.different_first_page_header_footer = True
    usable_width = section.page_width - section.left_margin - section.right_margin

    # running header (pages 2+): journal line ... "N of M"
    hdr_p = section.header.paragraphs[0]
    hdr_p.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
    r = hdr_p.add_run(JOURNAL_LINE); r.italic = True; r.font.size = Pt(8.5); r.font.name = BODY_FONT; r.font.color.rgb = MUTED
    hdr_p.add_run("\t")
    add_field(hdr_p, "PAGE"); hdr_p.add_run(" of "); add_field(hdr_p, "NUMPAGES")
    for run in hdr_p.runs:
        run.font.size = Pt(8.5); run.font.name = BODY_FONT; run.font.color.rgb = MUTED

    # footer (all pages): DOI line, right-aligned
    for footer, is_first in ((section.footer, False), (section.first_page_footer, True)):
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = fp.add_run(DOI_LINE); r.font.size = Pt(8); r.font.name = BODY_FONT; r.font.color.rgb = MUTED
    section.first_page_header.paragraphs[0].add_run("")  # explicitly blank on page 1

    # ---- parse front matter (Article / Title / Authors / Abstract / Keywords) --
    lines = SRC.read_text(encoding="utf-8").split("\n")
    n = len(lines)
    i = 0
    title_line, author_lines, abstract_buf, keywords_line = "", [], [], ""
    while i < n:
        s = lines[i].strip()
        if s == "Article":
            i += 1; continue
        if s.startswith("# "):
            title_line = s[2:]; i += 1; continue
        if s.startswith("**Abstract:**"):
            abstract_buf.append(s[len("**Abstract:**"):].strip()); i += 1
            while i < n and lines[i].strip() and not lines[i].strip().startswith("**Keywords"):
                abstract_buf.append(lines[i].strip()); i += 1
            continue
        if s.startswith("**Keywords:**"):
            keywords_line = s[len("**Keywords:**"):].strip(); i += 1; continue
        if s.startswith("## "):
            break
        if s and not s.startswith("---"):
            author_lines.append(s)
        i += 1

    # ---- render front matter ----------------------------------------------------
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Article"); r.italic = True; r.font.size = Pt(9); r.font.name = BODY_FONT; r.font.color.rgb = MUTED

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14); p.paragraph_format.line_spacing = 1.05
    add_inline_runs(p, title_line, base_size=19, base_bold=True, font=HEAD_FONT, color=INK)

    for idx, al in enumerate(author_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2 if idx == 0 else 1)
        is_author_name_line = idx == 0
        add_inline_runs(p, al, base_size=10.5 if is_author_name_line else 8.5,
                         base_italic=not is_author_name_line)
    if author_lines:
        doc.paragraphs[-1].paragraph_format.space_after = Pt(16)

    # sidebar (Received/Revised/Accepted/Published + copyright) beside Abstract+Keywords
    sidebar_table = doc.add_table(rows=1, cols=2)
    sidebar_table.autofit = False
    sidebar_table.columns[0].width = Cm(4.4)
    sidebar_table.columns[1].width = usable_width - Cm(4.4) - Cm(0.4)
    left_cell, right_cell = sidebar_table.rows[0].cells
    left_cell.width = Cm(4.4); right_cell.width = usable_width - Cm(4.4) - Cm(0.4)
    set_no_borders(sidebar_table)

    left_p = left_cell.paragraphs[0]
    for label in ("Received:", "Revised:", "Accepted:", "Published:"):
        pp = left_p if label == "Received:" else left_cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(2)
        add_inline_runs(pp, label, base_size=7.5, base_bold=True, color=MUTED)
    spacer = left_cell.add_paragraph(); spacer.paragraph_format.space_after = Pt(6)
    copyright_p = left_cell.add_paragraph()
    add_inline_runs(
        copyright_p,
        "Copyright: © 2026 by the authors. Submitted to *Journal Not Specified* for possible open "
        "access publication under the terms and conditions of the Creative Commons Attribution (CC BY) license.",
        base_size=7.5, color=MUTED,
    )

    right_p = right_cell.paragraphs[0]
    add_inline_runs(right_p, "Abstract", base_size=10.5, base_bold=True, font=HEAD_FONT, color=INK)
    abs_p = right_cell.add_paragraph(); abs_p.paragraph_format.space_after = Pt(10)
    abs_p.paragraph_format.line_spacing = 1.1
    add_inline_runs(abs_p, " ".join(abstract_buf), base_size=9.5)
    kw_p = right_cell.add_paragraph()
    add_inline_runs(kw_p, "Keywords: ", base_size=9.5, base_bold=True)
    add_inline_runs(kw_p, keywords_line, base_size=9.5, base_italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)  # breathing room before Section 1

    # ---- body: generic loop over everything from the first "## " heading -------
    # Numbered lists are rendered with a manually-typed "N." prefix rather than
    # Word's built-in "List Number" style: that style shares ONE numbering
    # instance document-wide by default, so a later independent list (e.g.
    # References) would silently continue counting from an earlier one (e.g.
    # the 3-item protocol list in Section 2.8) instead of restarting at 1.
    in_code, code_buf = False, []
    list_counter, prev_was_numbered = 0, False
    while i < n:
        line = lines[i]
        if line.strip().startswith("```"):
            if not in_code:
                in_code, code_buf = True, []
            else:
                in_code = False
                add_code_block(doc, code_buf)
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        s = line.strip()
        if s == "" or s == "---":
            i += 1; continue
        if s.startswith("## "):
            heading_paragraph(doc, s[3:], 13, space_before=20, space_after=6)
            list_counter, prev_was_numbered = 0, False; i += 1; continue
        if s.startswith("### "):
            heading_paragraph(doc, s[4:], 11, space_before=14, space_after=4)
            list_counter, prev_was_numbered = 0, False; i += 1; continue
        if s.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i]); i += 1
            add_table(doc, table_lines)
            list_counter, prev_was_numbered = 0, False; continue
        if re.match(r"^\d+\.\s", s):
            list_counter = list_counter + 1 if prev_was_numbered else 1
            prev_was_numbered = True
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.9)
            p.paragraph_format.first_line_indent = Cm(-0.9)
            p.paragraph_format.tab_stops.add_tab_stop(Cm(0.9))
            add_inline_runs(p, f"{list_counter}.\t", base_size=10)
            add_inline_runs(p, re.sub(r"^\d+\.\s", "", s), base_size=10)
            i += 1; continue
        if s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_inline_runs(p, s[2:], base_size=10.5)
            list_counter, prev_was_numbered = 0, False; i += 1; continue

        list_counter, prev_was_numbered = 0, False
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for key in FIGURE_FILES:
            if s.startswith(f"**{key}**"):
                add_figure(doc, key)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline_runs(p, s, base_size=9, base_italic=False)
                p.paragraph_format.space_after = Pt(14)
                i += 1
                break
        else:
            is_note = s.startswith("*[")
            is_back_matter = any(s.startswith(f"**{lbl}**") for lbl in BACK_MATTER_LABELS)
            add_inline_runs(p, s, base_size=9.5 if is_note else 10.5)
            p.paragraph_format.space_after = Pt(10 if is_back_matter else 8)
            i += 1

    doc.save(str(DOCX_OUT))
    print("saved", DOCX_OUT)


# --------------------------------------------------------------------------- HTML
def build_html():
    def slugify(text):
        text = re.sub(r"<[^>]+>", "", text)
        text = re.sub(r"[^\w\s-]", "", text).strip().lower()
        return re.sub(r"\s+", "-", text)

    def inline(text):
        text = protect_escapes(text)
        text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
        text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<em>\1</em>", text)
        text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
        text = re.sub(r"(?<![\"'>])(https?://[^\s<]+)", r'<a href="\1">\1</a>', text)
        return restore_escapes(text)

    def data_uri(fname):
        b64 = base64.b64encode((FIGDIR / fname).read_bytes()).decode("ascii")
        return f"data:image/png;base64,{b64}"

    def parse_table(tlines):
        rows = [l for l in tlines if l.strip().startswith("|")]
        data = []
        for idx, r in enumerate(rows):
            cells = [c.strip() for c in r.strip().strip("|").split("|")]
            if idx == 1 and all(re.match(r"^:?-{2,}:?$", c) for c in cells):
                continue
            data.append(cells)
        if not data:
            return ""
        head, body = data[0], data[1:]
        out = ['<div class="tw"><table><thead><tr>']
        out += [f"<th>{inline(c)}</th>" for c in head]
        out.append("</tr></thead><tbody>")
        for row in body:
            out.append("<tr>")
            for c in row:
                numeric = bool(re.match(r"^-?[\d.]+%?$|^\d\.\d\d$", c.strip()))
                out.append(f'<td class="num">{inline(c)}</td>' if numeric else f"<td>{inline(c)}</td>")
            out.append("</tr>")
        out.append("</tbody></table></div>")
        return "\n".join(out)

    lines = SRC.read_text(encoding="utf-8").split("\n")
    n, i = len(lines), 0
    body_parts, toc = [], []
    abstract_buf, keywords_line, title_line = [], "", ""

    while i < n:
        s = lines[i].strip()
        if s.startswith("# "):
            title_line = s[2:]; i += 1; continue
        if s.startswith("**Abstract:**"):
            abstract_buf.append(s[len("**Abstract:**"):].strip()); i += 1
            while i < n and lines[i].strip() and not lines[i].strip().startswith("**Keywords"):
                abstract_buf.append(lines[i].strip()); i += 1
            continue
        if s.startswith("**Keywords:**"):
            keywords_line = s[len("**Keywords:**"):].strip(); i += 1; continue
        if s.startswith("## "):
            break
        i += 1

    author_block = (
        '<p class="authors">[Author Name(s) to be completed]<sup>1,*</sup></p>'
        '<p class="affil"><sup>1</sup> [Affiliation to be completed]</p>'
        '<p class="corr">* Correspondence: <a href="mailto:ltarazonatorres@gmail.com">ltarazonatorres@gmail.com</a></p>'
    )

    in_code, code_buf = False, []
    while i < n:
        line = lines[i]; s = line.strip()
        if s.startswith("```"):
            if not in_code:
                in_code, code_buf = True, []
            else:
                in_code = False
                body_parts.append(f'<pre><code>{"".join(c + chr(10) for c in code_buf)}</code></pre>')
            i += 1; continue
        if in_code:
            code_buf.append(line.replace("<", "&lt;").replace(">", "&gt;")); i += 1; continue
        if s == "" or s == "---":
            i += 1; continue
        if s.startswith("## "):
            hid = slugify(s[3:]); body_parts.append(f'<h2 id="{hid}">{inline(s[3:])}</h2>'); toc.append((1, hid, s[3:])); i += 1; continue
        if s.startswith("### "):
            hid = slugify(s[4:]); body_parts.append(f'<h3 id="{hid}">{inline(s[4:])}</h3>'); toc.append((2, hid, s[4:])); i += 1; continue
        if s.startswith("|"):
            tlines = []
            while i < n and lines[i].strip().startswith("|"):
                tlines.append(lines[i]); i += 1
            body_parts.append(parse_table(tlines)); continue
        if re.match(r"^\d+\.\s", s):
            items = []
            while i < n and re.match(r"^\d+\.\s", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s", "", lines[i].strip())); i += 1
            body_parts.append("<ol>" + "".join(f"<li>{inline(it)}</li>" for it in items) + "</ol>"); continue
        if s.startswith("- "):
            items = []
            while i < n and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:]); i += 1
            body_parts.append("<ul>" + "".join(f"<li>{inline(it)}</li>" for it in items) + "</ul>"); continue
        matched_fig = next((f for k, f in FIGURE_FILES.items() if s.startswith(f"**{k}**")), None)
        if matched_fig:
            body_parts.append(f'<figure><img src="{data_uri(matched_fig)}" alt="{inline(s)}"><figcaption>{inline(s)}</figcaption></figure>')
            i += 1; continue
        cls = ' class="note"' if s.startswith("*[") else ""
        body_parts.append(f"<p{cls}>{inline(s)}</p>")
        i += 1

    abstract_html = inline(" ".join(abstract_buf))
    keywords_html = inline(keywords_line)
    body_html = "\n".join(body_parts)
    toc_html = "\n".join(f'<li class="lvl{lvl}"><a href="#{hid}">{inline(t)}</a></li>' for lvl, hid, t in toc)

    html = _HTML_TEMPLATE.format(
        title=inline(title_line), author_block=author_block,
        abstract=abstract_html, keywords=keywords_html, body=body_html, toc=toc_html,
    )
    HTML_OUT.write_text(html, encoding="utf-8")
    print("saved", HTML_OUT, len(html), "bytes")


_HTML_TEMPLATE = r"""<!doctype html>
<title>Ore-Hauling Simheuristics</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,500;9..144,600;9..144,700&family=Source+Serif+4:opsz,wght@8..60,400;8..60,500;8..60,600&family=IBM+Plex+Mono:wght@400;500;600&display=swap">
<style>
:root{{
  --paper:#EEF0EA; --surface:#F7F8F4; --surface-2:#E4E7DF; --ink:#20241D; --ink-2:#454B40;
  --muted:#6B7266; --hairline:#CBD1C3; --hairline-strong:#A9B29A; --accent:#2B4C63;
  --accent-2:#C97A3D; --accent-soft:#DCE4E3; --ok:#3C6E52; --warn:#8A5E0F;
}}
@media (prefers-color-scheme: dark){{
  :root:not([data-theme="light"]){{
    --paper:#15181A; --surface:#1B1F1E; --surface-2:#232823; --ink:#E7E9E1; --ink-2:#C2C8BB;
    --muted:#8D9587; --hairline:#2E342E; --hairline-strong:#454C41; --accent:#7FA7C4;
    --accent-2:#DB9A63; --accent-soft:#212C2E; --ok:#77B896; --warn:#D9A94E;
  }}
}}
:root[data-theme="dark"]{{
  --paper:#15181A; --surface:#1B1F1E; --surface-2:#232823; --ink:#E7E9E1; --ink-2:#C2C8BB;
  --muted:#8D9587; --hairline:#2E342E; --hairline-strong:#454C41; --accent:#7FA7C4;
  --accent-2:#DB9A63; --accent-soft:#212C2E; --ok:#77B896; --warn:#D9A94E;
}}
*{{box-sizing:border-box;}}
body{{margin:0;background:var(--paper);color:var(--ink);font-family:"Source Serif 4",Georgia,"Times New Roman",serif;font-size:17px;line-height:1.64;-webkit-font-smoothing:antialiased;}}
.page{{max-width:1200px;margin:0 auto;padding:0 24px 100px;}}
.grid{{display:grid;grid-template-columns:minmax(0,1fr);}}
@media (min-width:1040px){{.grid{{grid-template-columns:230px minmax(0,720px);column-gap:52px;justify-content:center;}}}}
.masthead{{grid-column:1/-1;border-bottom:2px solid var(--ink);padding:44px 0 26px;}}
.eyebrow{{font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:11px;font-weight:500;letter-spacing:.15em;text-transform:uppercase;color:var(--accent);margin:0 0 16px;}}
h1.title{{font-family:Fraunces,"Iowan Old Style",Georgia,serif;font-weight:600;font-size:clamp(28px,4vw,42px);line-height:1.12;letter-spacing:-.01em;margin:0 0 22px;text-wrap:balance;color:var(--ink);max-width:44ch;}}
.authors{{font-size:15px;color:var(--ink-2);margin:0 0 3px;}}
.affil,.corr{{font-size:12.5px;color:var(--muted);margin:0 0 3px;max-width:60ch;}}
.corr a{{color:var(--accent);}}
.meta-row{{display:flex;flex-wrap:wrap;gap:8px 26px;margin-top:20px;font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:11px;color:var(--muted);}}
.meta-row b{{color:var(--ink-2);font-weight:500;}}
nav.toc{{display:none;}}
@media (min-width:1040px){{
  nav.toc{{display:block;grid-column:1;position:sticky;top:0;align-self:start;padding-top:40px;max-height:100vh;overflow-y:auto;}}
  nav.toc ol{{list-style:none;margin:0;padding:0;}}
  nav.toc li.lvl2 a{{padding-left:22px;font-size:11px;}}
  nav.toc a{{display:block;font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:11.5px;line-height:1.4;color:var(--muted);text-decoration:none;padding:6px 0 6px 12px;border-left:2px solid var(--hairline);}}
  nav.toc a:hover,nav.toc a:focus-visible{{color:var(--accent);border-left-color:var(--accent);}}
}}
main{{grid-column:1;min-width:0;padding-top:36px;}}
@media (min-width:1040px){{main{{grid-column:2;}}}}
.abstract-block{{background:var(--surface);border:1px solid var(--hairline);border-left:3px solid var(--accent);padding:22px 26px 8px;margin:0 0 20px;}}
.abstract-block h2{{font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:11px;letter-spacing:.14em;text-transform:uppercase;color:var(--accent);border:none;margin:0 0 12px;padding:0;}}
.abstract-block p{{margin:0 0 14px;font-size:15.5px;}}
.keywords{{font-size:13px;color:var(--ink-2);margin:0 0 34px;}}
.keywords b{{font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:10.5px;letter-spacing:.1em;text-transform:uppercase;color:var(--muted);margin-right:8px;}}
h2{{font-family:Fraunces,Georgia,serif;font-weight:600;font-size:clamp(21px,2.6vw,26px);line-height:1.18;letter-spacing:-.01em;margin:52px 0 16px;padding-top:20px;border-top:1px solid var(--hairline-strong);text-wrap:balance;color:var(--ink);}}
main>h2:first-of-type{{margin-top:0;border-top:none;padding-top:0;}}
h3{{font-family:Fraunces,Georgia,serif;font-weight:600;font-size:18px;letter-spacing:-.008em;line-height:1.3;margin:30px 0 10px;color:var(--ink);}}
p{{margin:0 0 15px;max-width:70ch;}}
p.note{{font-size:14px;color:var(--muted);max-width:70ch;}}
ul,ol{{max-width:68ch;margin:0 0 16px;padding-left:22px;}}
li{{margin-bottom:6px;}}
strong{{font-weight:600;color:var(--ink);}}
a{{color:var(--accent);text-decoration-thickness:1px;text-underline-offset:2px;}}
.tw{{overflow-x:auto;margin:6px 0 22px;border:1px solid var(--hairline);background:var(--surface);}}
table{{border-collapse:collapse;width:100%;font-size:13px;}}
th,td{{text-align:left;padding:8px 12px;border-bottom:1px solid var(--hairline);vertical-align:top;}}
th{{font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:10px;font-weight:600;letter-spacing:.06em;text-transform:uppercase;color:var(--muted);background:var(--surface-2);white-space:nowrap;}}
td.num,th.num{{text-align:right;font-variant-numeric:tabular-nums;font-family:"IBM Plex Mono",ui-monospace,monospace;}}
tbody tr:last-child td{{border-bottom:none;}}
tbody tr:hover{{background:var(--surface-2);}}
figure{{margin:8px 0 26px;padding:0;}}
figure img{{max-width:100%;height:auto;border:1px solid var(--hairline);background:#fff;display:block;}}
figcaption{{font-size:13.5px;color:var(--ink-2);margin-top:10px;max-width:70ch;line-height:1.5;}}
figcaption strong{{color:var(--ink);}}
pre{{font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:12.5px;line-height:1.6;background:var(--surface);border:1px solid var(--hairline);border-left:3px solid var(--hairline-strong);padding:14px 16px;overflow-x:auto;margin:0 0 20px;white-space:pre;}}
code{{font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:.87em;}}
footer{{grid-column:1/-1;margin-top:70px;padding-top:20px;border-top:2px solid var(--ink);font-family:"IBM Plex Mono",ui-monospace,monospace;font-size:11px;color:var(--muted);line-height:1.7;}}
@media (prefers-reduced-motion:reduce){{*{{animation:none!important;transition:none!important;}}}}
</style>
<div class="page"><div class="grid">
<header class="masthead">
  <p class="eyebrow">Preprint &middot; Verification &amp; Validation Study &middot; Operations Research in Mining</p>
  <h1 class="title">{title}</h1>
  {author_block}
  <div class="meta-row"><span><b>Repository commit:</b> 4b3c2c0</span><span><b>Status:</b> Draft</span></div>
</header>
<nav class="toc" aria-label="Table of contents"><ol>{toc}</ol></nav>
<main>
<section class="abstract-block"><h2>Abstract</h2><p>{abstract}</p></section>
<p class="keywords"><b>Keywords</b>{keywords}</p>
{body}
</main>
<footer>Simheuristics for Stochastic Ore-Hauling Truck Dispatching &middot; ore_hauling_simheuristics<br>Regenerated by docs/build_paper.py from PAPER_ore_hauling_simheuristics.md</footer>
</div></div>
"""

if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else "all"
    if target in ("docx", "all"):
        build_docx()
    if target in ("html", "all"):
        build_html()
